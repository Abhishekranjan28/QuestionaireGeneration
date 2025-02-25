import streamlit as st
st.set_page_config(page_title="Chatbot-Report Generator", layout="wide")
from streamlit_cookies_manager import EncryptedCookieManager
import os
from dotenv import load_dotenv
from fpdf import FPDF
import tempfile
import fitz 
from docx import Document 
import google.generativeai as genai
import requests
from bs4 import BeautifulSoup
from PIL import Image
import psycopg2
import hashlib
import streamlit.components.v1 as components
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail
import random
from Automation import connect_db,insert_food,insert_clothes,insert_cosmetics

load_dotenv()

SENDGRID_API_KEY = os.environ.get("SENDGRID_API_KEY")
SEND_FROM_EMAIL = "corp@altibbe.com"

cookies = EncryptedCookieManager(
    prefix="my_app",
    password=os.getenv("COOKIES_PASSWORD", "your_secret_password"),
)

if not cookies.ready():
    st.stop()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    st.error("Gemini API key is not set in environment variables.")
    st.stop()


genai.configure(api_key=GEMINI_API_KEY)


DATABASE_URL = os.getenv("NEON_DB_URL1")
if not DATABASE_URL:
    st.error("PostgreSQL connection string is not set in environment variables.")
    st.stop()

def get_db_connection():
    try:
        conn = psycopg2.connect(DATABASE_URL, sslmode="require")
        return conn
    except Exception as e:
        st.error(f"Error connecting to the database: {e}")
        return None

def scrape_multiple_websites(urls):
    all_headings_and_articles = []
    for url in urls:
        try:
            response = requests.get(url)
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                heading = soup.find('h1').get_text(strip=True) if soup.find('h1') else "No heading found"
                article = " ".join([p.get_text(strip=True) for p in soup.find_all('p')])
                all_headings_and_articles.append((heading, article))
            else:
                st.error(f"Failed to fetch the URL: {url}. Status code: {response.status_code}")
                all_headings_and_articles.append(("Failed to fetch", ""))
        except Exception as e:
            st.error(f"An error occurred while scraping: {e}")
            all_headings_and_articles.append(("Error", ""))
    return all_headings_and_articles


def save_summary_to_db(table_name,summary):
    if table_name == "food":
        insert_food(summary)
    elif table_name == "clothes":
        insert_clothes(summary)
    elif table_name == "cosmetics":
        insert_cosmetics(summary)
    else:
        st.error("Invalid table selection.")

def fetch_summaries(table_name):
    """Fetch summaries from the selected table."""
    conn = connect_db()
    if conn:
        cursor = conn.cursor()
        try:
            cursor.execute(f"SELECT summary FROM {table_name}")
            rows = cursor.fetchall()
            if rows:
                additional_content = "\n".join([f"{idx}. {row[0]}" for idx, row in enumerate(rows, start=1)])
                return additional_content
            else:
                return "No summaries found in the database."
        except Exception as e:
            return f"Error retrieving summaries: {e}"
        finally:
            cursor.close()
            conn.close()
    return "Database connection failed."

def initialize_db():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            email TEXT PRIMARY KEY,
            password TEXT NOT NULL
        )
    ''')
    conn.commit()
    cursor.close()
    conn.close()

initialize_db()  

def generate_otp():
    return str(random.randint(100000, 999999))

def send_email(otp, recipient_email):
    message = Mail(
        from_email=SEND_FROM_EMAIL,
        to_emails=recipient_email,
        subject="Your OTP for Email Verification",
        html_content=f"<p>Your OTP is: <strong>{otp}</strong></p>")
    
    try:
        sg = SendGridAPIClient(SENDGRID_API_KEY)
        response = sg.send(message)
        return response.status_code
    except Exception as e:
        return str(e)

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def register_user(email, password):
    conn = get_db_connection()
    cursor = conn.cursor()
    hashed_pw = hash_password(password)
    cursor.execute("INSERT INTO users (email, password) VALUES (%s, %s) ON CONFLICT (email) DO NOTHING", (email, hashed_pw))
    conn.commit()
    cursor.close()
    conn.close()

def login_user(email, password):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT password FROM users WHERE email=%s", (email,))
    result = cursor.fetchone()
    conn.close()
    return result and result[0] == hash_password(password)

if "logged_in" not in st.session_state:
    st.session_state.logged_in = cookies.get("logged_in") == "true"
    st.session_state.user_email = cookies.get("user_email")

if "otp" not in st.session_state:
    st.session_state.otp = None
if "temp_email" not in st.session_state:
    st.session_state.temp_email = None
if "temp_password" not in st.session_state:
    st.session_state.temp_password = None

if not st.session_state.logged_in:

    col_left, _, col_right = st.columns([1, 6, 1]) 
    with col_left:
        st.image("Altibbe logo dark.png", width=150)

    with col_right:
        st.image("Hedamo.jpg", width=200)

    st.markdown("<h1 style='font-size:24px; color:black;'>Login and Register</h1>", unsafe_allow_html=True)

    auth_mode = st.radio("Choose Authentication Mode", ["Login (If registered)", "Register"])
    email = st.text_input("Enter your email:")
    password = st.text_input("Enter your password:", type="password")

    if auth_mode == "Register":
        if st.button("Send OTP"):
            if email and password:
                otp = generate_otp()
                st.session_state.otp = otp
                st.session_state.temp_email = email
                st.session_state.temp_password = password
                status = send_email(otp, email)
                if status == 202:
                    st.success(f"OTP sent successfully to {email}! ✅")
                else:
                    st.error(f"Failed to send OTP: {status}")
            else:
                st.warning("Please enter both email and password.")

        otp_input = st.text_input("Enter the OTP received:")
        if st.button("Verify OTP"):
            if otp_input and st.session_state.otp and otp_input == st.session_state.otp:
                register_user(st.session_state.temp_email, st.session_state.temp_password)
                st.success(f"Email {st.session_state.temp_email} verified and registered successfully! ✅")
                st.session_state.otp = None
                st.session_state.temp_email = None
                st.session_state.temp_password = None
            else:
                st.error("Invalid OTP. Please try again.")

    elif auth_mode == "Login (If registered)":
        if st.button("Login"):
            if email and password:
                if login_user(email, password): 
                    st.session_state.logged_in = True
                    st.session_state.user_email = email
                    st.success("Logged in successfully! 🎉")
                    components.html(
                        f"""
                        <script>
                            document.cookie = "logged_in=true; path=/";
                            document.cookie = "user_email={email}; path=/";
                        </script>
                        """,
                        height=0
                    )
                    st.rerun()  
                else:
                    st.error("Invalid email or password. Please try again.")

else:
    if st.button("Logout"):
        cookies["logged_in"] = "false"
        cookies["user_email"] = ""
        cookies.save()
        st.session_state.logged_in = False
        st.session_state.user_email = None
        st.rerun()   

    col_left, _, col_right = st.columns([1, 6, 1]) 

    with col_left:
        st.image("Altibbe logo dark.png", width=150)

    with col_right:
        st.image("Hedamo.jpg", width=200)

    st.markdown(
    """
    <a href="https://huggingface.co/spaces/abhi280622/Bhai-AI" target="_blank">
        <button style="padding:8px 15px; font-size:16px;">Go to Report Generator Chatbot</button>
    </a>
    """,
    unsafe_allow_html=True
    )
    
    st.markdown(
    """
    <a href="https://huggingface.co/spaces/random29/QuestionaireGeneratorViaScrapping" target="_blank">
        <button style="padding:8px 15px; font-size:16px;">Scrapper Questionaire Generator</button>
    </a>
    """,
    unsafe_allow_html=True
    )

    if st.session_state.get('allow_admin_download'):
                   st.subheader("Admin Downloads")
                   st.download_button("Download CSV of Generated Questions", data=open("generated_questions.csv", "rb"), file_name="generated_questions.csv", mime="text/csv")
                   st.download_button("Download CSV of Registered Users", data=open("users_data.csv", "rb"), file_name="users_data.csv", mime="text/csv")
    st.markdown(
        """
        <h1 style='font-size:24px; color:black;'>Questionnaire Generator</h1>
        """,
        unsafe_allow_html=True,
    )

    st.write("Generate questions based on input data, including PDFs, text files, or Word documents.")

    st.markdown(
    """
    <h1 style='font-size:20px; color:black;'>Upload Files</h1>
    """, 
    unsafe_allow_html=True
)
    
    extracted_text = ""
    uploaded_files = st.file_uploader(
        "Upload files (PDF, text, or Word documents):",
        type=["pdf", "txt", "docx"],
        accept_multiple_files=True,
    )

    if uploaded_files:
          for uploaded_file in uploaded_files:
            file_extension = uploaded_file.name.split(".")[-1].lower()
            try:
                if file_extension == "pdf":
                    with fitz.open(stream=uploaded_file.read(), filetype="pdf") as pdf_file:
                        for page_num in range(pdf_file.page_count):
                            page = pdf_file[page_num]
                            extracted_text += page.get_text("text") + "\n"
                elif file_extension == "txt":
                    extracted_text += uploaded_file.read().decode("utf-8") + "\n"
                elif file_extension == "docx":
                    doc = Document(uploaded_file)
                    for para in doc.paragraphs:
                        extracted_text += para.text + "\n"
            except Exception as e:
                st.error(f"Error processing file '{uploaded_file.name}': {str(e)}")
     
    st.markdown(
    """
    <h1 style='font-size:20px; color:black;'>Additional Information</h1>
    """, 
    unsafe_allow_html=True
)
    specific_constraints = st.text_area(
        "Enter specific constraints or additional information for the questionnaire generation (optional):",
        height=150,
    )

    st.markdown(
    """
    <h1 style='font-size:20px; color:black;'>Product Details</h1>
    """, 
    unsafe_allow_html=True
)
    
    combined_article=""

    website_urls_input = st.text_area("Enter multiple website URLs (comma separated):(optional)")

    category = st.selectbox("Select the product category:", ["food", "clothes", "cosmetics"])

    st.markdown(
    """
    <h1 style='font-size:20px; color:black;'>Upload Image</h1>
    """, 
    unsafe_allow_html=True
)
    uploaded_images = st.file_uploader(
    "Upload images related to the product or company (you can select multiple):",
    type=["jpg", "png", "jpeg"],
    accept_multiple_files=True,
)

    image_context = ""

    st.markdown(
    """
    <h1 style='font-size:20px; color:black;'>Questionaire Generator Settings</h1>
    """, 
    unsafe_allow_html=True
)
    num_questions = st.number_input(
        "Enter the number of questions to generate:",
        min_value=1,
        max_value=200,
        value=30,
    )
    
    def generate_pdf(questions):
       pdf = FPDF()
       pdf.add_page()

       pdf.add_font('Roboto', '', 'Roboto-Regular.ttf', uni=True)
       pdf.add_font('Roboto-Bold', 'B', 'Roboto-Bold.ttf', uni=True)
       pdf.set_font('Roboto', size=12)

       pdf.cell(200, 10, txt="Generated Questions", ln=True, align='C')
       pdf.ln(10)

       for index, question in enumerate(questions, start=1):  
          pdf.multi_cell(0, 10, f"{question}")

       with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
           pdf_output_path = tmp_file.name
           pdf.output(pdf_output_path)

       return pdf_output_path

    if st.button("Generate Deep Questioning Questions"):
       
            try:
                if website_urls_input:
                       urls = [url.strip() for url in website_urls_input.split(",")]
    
                       scraped_data = scrape_multiple_websites(urls)

                       st.markdown("<h5>Scraped Data:</h5>", unsafe_allow_html=True)
                       for idx, (heading, article) in enumerate(scraped_data):
                          st.markdown(f"##### Website {idx + 1}: {heading}")
                          st.write(article)

                       combined_article = "\n".join([article for _, article in scraped_data])
                       st.markdown(f"##### Combined Article from All Websites:")
                       st.write(combined_article)

                if uploaded_images:
                  image_paths = []
                  st.markdown("<h4>Uploaded Images:</h4>", unsafe_allow_html=True)
                  for idx, uploaded_image in enumerate(uploaded_images, start=1):

                    image = Image.open(uploaded_image)
                    st.image(image, caption=f"Image {idx}", use_container_width=True)
        
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
                       image.save(temp_file.name)
                       image_paths.append(temp_file.name)
    
    
                  image_context = (
                    "Images related to the product/company have been uploaded. File paths:\n" +
                    "\n".join(image_paths)
                  )
                
                summaries = fetch_summaries(category)
    
                prompt = (
        f"The questionnaire should focus on key areas relevant to health, quality, safety, sustainability, and brand positioning, "
        f"specifically addressing the unique health benefits and consumer perceptions surrounding the product, with a strong emphasis on potential hazards related to agricultural inputs, post-harvest treatments, seed quality, and ripening processes.\n\n"
        f"***Context from input files:***\n{extracted_text}\n\n + {combined_article}\n\n +{image_context}\n\n"
        f"***Highest priority should be given to the context from input files.***\n\n"
        f"**Specific Constraints or Information:**\n{specific_constraints}\n\n"
        f"**Instructions:**\n"
        f"- Generate {2*num_questions // 3} Text-based responses type questions and {num_questions //3} Text + File Upload responses type questions. Divide the Questions in Two Sections 1. Text based Responses 2.Text and file upload responses.\n\n"
        f"- **Tone of Questions should be Polite and Professional such as Questions must start with ***Please or Kindly***.Must avoid using ***Can/could*** as starting words of Questions.** Follow instructions strictly.\n\n"
    f"**Key Objectives:**\n"
    f"1. Questions that require certificates or supporting documents to explicitly ask for file uploads.\n\n"
    f"2. Investigate product-specific attributes such as health benefits, safety, and unique health propositions, with a focus on potential contamination and chemical residues originating from the specified inputs and treatments.\n\n"
    f"3. Assess good practices, compliance with regulations, certifications, and ethical practices in business, specifically related to the use of agricultural inputs, post-harvest treatments, seed sourcing, and ripening methods, paying close attention to the potential health and environmental impacts described in the provided text.\n\n"
    f"4. Understand the current and planned market presence, brand perception, and pricing strategy, considering consumer concerns about food safety and environmental impact, particularly concerning the issues highlighted in the provided text.\n\n"
    f"5. Inquire about production challenges, competitor analysis, and growth-oriented goals, with an emphasis on mitigating risks related to contamination and chemical residues, drawing on the information provided about specific hazards.\n\n"
    f"6. Gather detailed insights into sustainable practices, including resource efficiency, waste management, and the reduction of harmful chemical usage, with reference to the specific chemicals and practices mentioned in the text.\n\n"
    f"7. Capture information regarding the future of company and its commitments to health and sustainability improvements, particularly in relation to minimizing or eliminating the use of harmful substances identified in the provided text.\n\n"
    f"8. Generate Questions in form such that my client has to answer about his product.\n"
    f"9. Don't Use symbols of currency. Rather use Name of Currency in response.\n"
    f"10. Don't use some special symbols(like smart apostrophe) that cannot be encoded using codec.\n"
    f"11. Use good Introduction and conclusion.\n"
    f"12. Question should be of top Quality.\n"
    f"13. Include scientific, geographic, different stages of production process with more emphasis while generating questions.\n"
    f"14. Questions should be detailed and interesting to answer.\n"
    f"15. **Put all Questions of similar topic at once under a sub-heading and do not include question topic separately with each questions.**\n"
    f"16. Generate questions on different processes of making the product.Ex: Different stages of crop production,processing in industries etc.\n"
    f"17. Also add the topic along with the generated questions( examples like: Health Attribute, production type, sustainability, marketing etc.)\n"
    f"18. Questions should be in minimum 40 words and very descriptive and explained.\n\n"
    f"19. Generate questions in such a way that We are assuming that the producers is all clear, but we are going to question them deeply to remove all doubt. Don't start questions with validatory terms.\n"
    f"20. Be very strict to given instructions.Generate exact number of questions given in instructions.\n"
    f"- Ensure questions are focused on product quality, compliance, good practices, brand positioning, and sustainability, with a strong emphasis on the following areas, drawing specific details from the provided text:\n"
    f"  **Harmful agricultural inputs:** (Pesticides: Organophosphates, Carbamates, Synthetic Pyrethroids, Organochlorines, Neonicotinoids; Herbicides: Glyphosate; Fungicides: Mancozeb; Fertilizers: Synthetic Nitrogen, Phosphorus, Potassium; PGRs: Synthetic Hormones, Non-compliant Auxins/Cytokinins; Soil Conditioners: Heavy Metals, Sewage Sludge; Animal Feed Additives: Antibiotics, Hormonal Additives; Contaminated Irrigation Water: Industrial Chemicals, Pesticide Runoff)\n"
    f"  **Post-harvest treatments:** (Synthetic Preservatives: Sulfur Dioxide, Benzoates, Sorbates; Wax Coatings: Natural/Synthetic; Synthetic Fumigants: Methyl Bromide, Phosphine Gas; Chemical Treatments for Disease Prevention: Imazalil, Thiabendazole, Chlorine Washes; Cold Storage/MAP: Chemical Leaching, Nutrient Loss; Irradiation: Radiolytic Products, Nutrient Loss; Hazardous Packaging: Plasticizers, BPA)\n"
    f"  **Contaminated seeds:** (Chemical Treatments: Carbendazim, Captan, Thiram; Pathogenic Contamination: Fungi, Bacteria, Viruses; Environmental Exposure: Industrial Pollutants, Heavy Metals; Improper Storage/Handling; Adulteration)\n"
    f"  **Chemical ripening agents:** (Calcium Carbide: Arsenic, Phosphorus; Ethylene Gas; Ethephon; Methyl Jasmonate)\n\n"
    f"  **Harmful Effects of Additives and Antibiotics in Livestock Production on Human Health:(If the production type is of Livestock then include it otherwise don't)** This document outlines the risks associated with the use of hormonal additives (estrogen, testosterone, progesterone, growth-promoting steroids) and antibiotics in livestock feed. Hormonal additives can cause endocrine disruption (hormonal imbalances, reproductive issues, developmental abnormalities, increased risk of hormone-dependent cancers), cardiovascular risks, obesity, metabolic disorders, and neurological disorders. Environmental contamination from these additives can further impact human health through water and soil contamination and bioaccumulation. Antibiotic use contributes to antibiotic resistance, allergic reactions, disruption of gut microbiota, and potential carcinogenic and toxic effects. Mitigation strategies include improved animal husbandry, use of probiotics, prebiotics, and phytogenic feed additives, stricter regulations, and increased consumer awareness.\n\n"
    f"  **Harmful Effects of Contaminated Irrigation Water Due to Harmful Agricultural Practices**: This document highlights the detrimental impact of harmful agricultural practices on irrigation water quality, posing risks to human health and the environment. Key contaminants include pesticide runoff (Organochlorines, Organophosphates, Carbamates, Neonicotinoids), industrial chemicals (Arsenic, Mercury, Cadmium, Lead), fertilizer leaching (Nitrates), animal waste (Pathogens, Hormonal Residues), contaminated seeds (Thiram, Captan), and plastic mulches (Microplastics). These contaminants can cause neurological disorders, endocrine disruption, gastrointestinal and respiratory issues, carcinogenic risks, and chronic diseases in humans. Environmentally, they contribute to eutrophication, biodiversity loss, soil degradation, and bioaccumulation. Mitigation strategies include sustainable farming practices (IPM, organic farming), advanced irrigation techniques (drip irrigation, water filtration), proper waste management, monitoring and regulation, and education and awareness programs.\n\n"
    f"  **Adulterated Organic Manure Concerns:** Adulterated organic manures, containing impurities like heavy metals and industrial waste, negatively impact soil health (contamination, reduced microbial activity, altered pH), crop growth (toxicity, poor yield, residual effects), human health (food safety, pathogen transmission, chemical exposure), and the environment (water and air pollution, soil erosion). This leads to economic losses and erodes trust in organic practices. Stringent quality control, certification, and farmer awareness are crucial.\n\n"
    f"  **Harmful Effects of Chemical Residues on Human Health:** Chemical residues from pesticides (insecticides, fungicides, rodenticides) pose significant health risks. Acute exposure can cause neurological effects (headaches, confusion), respiratory issues (breathing difficulty), gastrointestinal problems (nausea, vomiting), and skin/eye irritation. Chronic exposure is linked to carcinogenic problems (e.g., DDT and cancer), endocrine disruption (e.g., chlorpyrifos and thyroid imbalances), neurotoxicity (e.g., organophosphates and memory loss), reproductive/developmental effects (e.g., dioxins and birth defects), and immunotoxicity (e.g., lindane and decreased white blood cell count). Residues are found in fruits, vegetables, meat, dairy, and water. Infants, children, pregnant women, farmers, and the elderly are particularly vulnerable. Prevention and mitigation strategies include regulatory measures (MRLs, bans), organic/sustainable farming (IPM), consumer awareness (washing produce, choosing organic), monitoring/testing, and education/training.\n\n"
    f"  **Harmful Effects of Plastics in Agriculture and the Food Industry:** The extensive use of plastics in agriculture and the food industry leads to significant environmental and health risks. Environmentally, plastics cause soil pollution (reduced aeration, water infiltration, and microbial activity), non-biodegradability (landfill accumulation), water pollution (microplastic contamination), and air pollution (from burning). In agriculture, plastics contribute to microplastic accumulation in soil, reduced soil productivity, and irrigation system blockages. In the food industry, they cause food contamination (chemical leaching of BPA and phthalates), and increased waste generation. Health impacts include chemical exposure and toxicity (endocrine disruption from BPA and phthalates, potential carcinogenicity from styrene), microplastic ingestion (inflammation, toxicity from attached chemicals, cellular damage), inhalation of airborne plastic particles (respiratory issues, neurological and developmental effects), and bioaccumulation in seafood (liver damage, neurological issues). Plastics also contribute to climate change through greenhouse gas emissions and fossil fuel dependence. Mitigation strategies include biodegradable alternatives, plastic recycling and reuse, government regulations, awareness campaigns, and adopting a circular economy.\n\n"
    f"  **Adulterants in Milk (Only when questions are asked about Diary products):** Milk is susceptible to adulteration with substances like water (dilutes nutrients, introduces microbes), starch (indigestion), detergents (gastrointestinal irritation, cancer risk), urea (kidney damage), formalin (toxic, carcinogenic), ammonium sulfate (organ damage), sodium carbonate/bicarbonate (digestive problems), synthetic milk (toxic residues), glucose/sugar (obesity, diabetes, cavities), hydrogen peroxide (toxic, cellular damage), and neutralizers (gastrointestinal irritation). These adulterants are added for economic gain or to increase shelf life. Detection methods include chemical tests, advanced techniques (chromatography, spectroscopy), and milk analyzers. Prevention relies on regulatory oversight, consumer awareness, and technological innovations. This emphasizes the need to scrutinize agricultural inputs, post-harvest treatments, seed quality, and all stages of the production process for potential contamination and adulteration.\n\n"
    f"  **If given products is about **Nuts** then include this as context**: Nuts are frequently adulterated to increase profit or shelf life with substances like artificial colorants (potentially causing allergic reactions, hyperactivity, and cancer), wax coatings (leading to digestive issues and accumulation of harmful substances), added salt and sugar (contributing to high blood pressure, obesity, and diabetes), artificial sweeteners (causing gastrointestinal problems, allergic reactions, headaches, and metabolic disturbances), cheap seeds/grains (reducing nutritional value and introducing contaminants), aflatoxins (mold toxins linked to liver cancer and other health issues), hydrogenated oils (containing trans fats that increase heart disease risk), sodium nitrite (potentially forming carcinogenic nitrosamines), formalin (causing severe health problems including cancer), and synthetic preservatives (potentially causing allergic reactions and damaging detoxification mechanisms); detection methods include visual inspection, lab testing, taste tests, and microscopic examination; prevention involves trusted sourcing, consumer awareness, regulatory oversight, and proper storage.\n\n"
    f"  **If given products is about **rice** then include this as context**: Rice is susceptible to adulteration impacting quality, safety, and nutrition via artificial colorants (potentially causing toxicity, allergies, cancer), starch additives (leading to nutritional degradation, digestive issues, chemical contamination), excessive polishing (resulting in nutrient loss and higher glycemic index), artificial fragrances (potentially causing respiratory issues, allergies, toxicity), foreign particles (causing physical injury, chemical contamination, illness), and arsenic contamination (increasing cancer and neurological/developmental risks); prevention involves organic farming, consumer awareness (labeling, trusted sources, washing/soaking), government regulations (pesticides, heavy metals), and improved processing (filtration, less polishing).\n\n"
    f"  **If given product is about **eggs** then include this as context** :Eggs are susceptible to adulteration with substances like water (increasing weight/contamination risk), artificial coloring (toxicity/carcinogenicity), malachite green (carcinogen/organ toxicity), antibiotics (resistance/hormonal disruption), formaldehyde (carcinogen/toxicity), hydrogen peroxide (toxicity), and shell coatings (contamination/masking defects), posing health risks like cancer, antibiotic resistance, and allergic reactions; detection involves testing, inspection, and smell, while prevention relies on trusted sourcing, regulation, and awareness.\n\n"
    f"  **If given product ia about **honey** then include this as context**: Proper honey storage involves maintaining low moisture (below 18%), temperatures between 50-77°F (avoiding >104°F and extreme cold), minimal light exposure (using opaque containers), airtight containers to prevent oxidation and contamination, and using glass or food-grade plastic containers (avoiding reactive metals and porous materials); improper storage can lead to fermentation, nutrient loss, contamination, and crystallization (reversible by gentle warming); best practices include airtight containers, cool/dry/dark storage, hygiene, and labeling, allowing for an indefinite shelf life when done correctly.\n\n"
    f"  **If given product is about **Fish** then include this as context** :Fish adulteration with chemicals (formalin, sodium benzoate, dyes), water/chemical injections, pesticides, heavy metals, phosphates, and antibiotics poses serious health risks like cancer, neurological damage, allergic reactions, and organ damage, requiring consumer vigilance, trusted sourcing, and strong regulatory measures.\n\n"
    f"  **If given product is about **Poultry Products like Chicken** then include this as context** :Malpractices across the poultry industry, from breeding and feeding to processing and marketing, involving unethical practices like substandard feed, hormone/antibiotic use, overcrowding, chemical contamination, mislabeling, and improper waste disposal, compromise product quality, endanger consumer health through antibiotic resistance, food poisoning, and chronic illnesses, and harm the environment, necessitating stronger regulations, education, traceability, sustainable practices, and rigorous testing. **Improper Storage in Poultry products like Chicken**Improper storage of poultry meat and eggs, including inadequate humidity/ventilation, cross-contamination, temperature abuse (freezer burn, improper refrigeration), chemical preservatives (formalin), dirty/cracked shells, moisture loss, and temperature fluctuations, leads to loss of nutritional quality, microbial growth (Salmonella, E. coli, Listeria, Campylobacter), spoilage, increased food waste, and human health impacts like chemical toxicity, foodborne illnesses (salmonellosis, listeriosis, campylobacteriosis), antibiotic resistance, and allergic reactions, necessitating consumer awareness, proper refrigeration, regular monitoring, and hygienic packaging.\n\n"
    f"  **If given product is about Poultry Products (meat and eggs) then include this as context**: Malpractices in poultry product transportation, including improper cold chain maintenance (temperature abuse, power failures, delays), cross-contamination (contaminated containers, improper segregation, packaging failures), and mishandling (breakage, inadequate containers, overcrowding), compromise product quality and safety, leading to spoilage, microbial contamination (Salmonella, Listeria), and foodborne illnesses (salmonellosis, listeriosis). Solutions involve hygienic transport practices, cold chain logistics, handling guidelines, regulatory compliance, and sustainable measures to minimize risks and ensure consumer health and safety.\n\n"
    f"  **If given product is about Poultry Products like Eggs then include this as context**: Forced molting in poultry, a practice inducing egg-laying cessation through feed restriction, light manipulation, or nutritional modification, causes significant animal welfare concerns including hunger, stress,  physical suffering, high mortality, immune suppression, and bone health issues.  This clashes with natural molting cycles, raising ethical contradictions in animal husbandry. While unregulated in many areas due to economic benefits, growing consumer awareness and stricter regulations in some regions are driving a shift towards welfare-friendly alternatives like selective breeding, natural molting, and improved management practices.\n\n"
    f"  **If given product is about Poultry Products like Chicken and Eggs then include this as context**:  High cholesterol in poultry, stemming from biological factors (species, organs), feeding practices (rich-fat diets, imbalances), farming practices (lack of exercise, selective breeding), and processing/cooking methods (deep frying, skin retention), poses human health risks (cardiovascular disease, obesity, chronic illnesses) and economic challenges for the poultry industry.  Mitigation strategies include nutritional interventions (low-fat diets, additives, omega-3s, enzyme supplementation), sustainable farming (selective breeding, free-range systems), post-harvest measures (discarding fat/skin, healthy cooking), genetic approaches (gene editing, modification), consumer education (lean cuts, cooking methods), and regulatory measures (feed quality monitoring, certifications).\n\n"
    f"  **If given product is about Mushrooms then include this as context**: Mushroom quality is compromised by microbial and fungal contamination, physical damage (bruising, cap damage), improper temperature and humidity control, short shelf life leading to nutrient loss and color/texture changes, pest infestations (insects, mites), pesticide residues, allergenic reactions (spores, proteins), toxic wild mushroom contamination, and poor water/substrate quality, necessitating improved cultivation, handling, storage, and regulatory compliance for safe, high-quality products.\n\n"
    f"  **If given product is about Dates then include this as context**: Unethical weight manipulation techniques in the date industry, including sugar syrup dipping, saltwater injection, excessive rehydration, improper drying, and addition of glycerin, polyethylene glycol, sand, stone powder, or starch, increase weight but compromise quality, causing health risks like toxicity, microbial contamination (aflatoxins), nutritional imbalance, and consumer deception.  Stringent quality control, consumer education, clear labeling, and legal enforcement are crucial to ensure safe and high-quality dates.\n\n"
    f"  **If given product is about Wax Coatings in the Date Industry then include this as context**:  The use of wax coatings (natural like beeswax, carnauba; synthetic like polyethylene) on dates enhances appearance and shelf life but poses health risks (chemical residues, digestive issues, toxicity), environmental concerns (non-biodegradable waxes, petroleum use), and regulatory challenges (lack of consistent standards).  Alternatives like edible coatings (aloe vera, starch, pectin) and bio-based polymers (chitosan) are needed to mitigate these negative impacts.\n\n"
    f"{summaries}"
    )           
                st.write(prompt)
                def generate_docx(questions):
                   """
                   Generates a DOCX file with questions and inputs and saves it to disk.
                   """
                   doc = Document()
                   doc.add_heading("Generated Questionnaire", level=1)
    
                   doc.add_heading("Generated Questions:", level=2)
                   for question in questions:
                      doc.add_paragraph(f"- {question}")
    
            
                   doc_path = "questionnaire_report.docx"
                   doc.save(doc_path)
                   return doc_path

                model = genai.GenerativeModel("gemini-2.0-flash-exp")
                response = model.generate_content(prompt)
                questions = response.text.strip().split("\n")

                st.subheader("Generated Questions:")
                for q in questions:
                    st.write(f"- {q}")
                
                pdf_path = generate_pdf(questions)
                docx_path = generate_docx(questions)
              
                with open(pdf_path, "rb") as pdf_file:
                     st.download_button(
                     "Download Full Report as PDF",
                     pdf_file,
                     file_name="questionnaire_report.pdf",
                    )
                with open(docx_path, "rb") as docx_file:
                     st.download_button(
                   "Download Full Report as DOCX",
                   docx_file,
                   file_name="questionnaire_report.docx",
        )
              
            except Exception as e:
                st.error(f"Error generating questions: {e}")

    if st.button("Generate Initial Assessment Questions"):
            try:
                if website_urls_input:
                       urls = [url.strip() for url in website_urls_input.split(",")]
    
                       scraped_data = scrape_multiple_websites(urls)

                       st.markdown("<h5>Scraped Data:</h5>", unsafe_allow_html=True)
                       for idx, (heading, article) in enumerate(scraped_data):
                          st.markdown(f"##### Website {idx + 1}: {heading}")
                          st.write(article)

                       combined_article = "\n".join([article for _, article in scraped_data])
                       st.markdown(f"##### Combined Article from All Websites:")
                       st.write(combined_article)

                if uploaded_images:
                  image_paths = []
                  st.markdown("<h4>Uploaded Images:</h4>", unsafe_allow_html=True)
                  for idx, uploaded_image in enumerate(uploaded_images, start=1):

                    image = Image.open(uploaded_image)
                    st.image(image, caption=f"Image {idx}", use_container_width=True)
        
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
                       image.save(temp_file.name)
                       image_paths.append(temp_file.name)
    
    
                  image_context = (
                    "Images related to the product/company have been uploaded. File paths:\n" +
                    "\n".join(image_paths)
                  )
                
                Text=""
                with open("text.txt", "r", encoding="utf-8") as file:
                    Text = [line.strip().replace('\uf0a7', '-') for line in file]
    
                prompt = (
        f"The questionnaire should focus on key areas relevant to health, quality, safety, sustainability, and brand positioning, "
        f"specifically addressing the unique health benefits and consumer perceptions surrounding the product, with a strong emphasis on potential hazards related to agricultural inputs, post-harvest treatments, seed quality, and ripening processes.\n\n"
        f"***Context from input files:***\n{extracted_text}\n\n + {combined_article}\n\n +{image_context}\n\n"
        f"***Highest priority should be given to the context from input files.***\n\n"
        f"**Specific Constraints or Information:**\n{specific_constraints}\n\n"
        f"**Instructions:**\n"
        f"- **Tone of Questions should be Polite and Professional such as Questions must start with ***Please or Kindly***.Must avoid using ***Can/could*** as starting words of Questions.** Follow instructions strictly.\n\n"
    f"**Key Objectives:**\n"
    f"1. Questions that require certificates or supporting documents to explicitly ask for file uploads.\n\n"
    f"2. Investigate product-specific attributes such as health benefits, safety, and unique health propositions, with a focus on potential contamination and chemical residues originating from the specified inputs and treatments.\n\n"
    f"3. Assess good practices, compliance with regulations, certifications, and ethical practices in business, specifically related to the use of agricultural inputs, post-harvest treatments, seed sourcing, and ripening methods, paying close attention to the potential health and environmental impacts described in the provided text.\n\n"
    f"4. Understand the current and planned market presence, brand perception, and pricing strategy, considering consumer concerns about food safety and environmental impact, particularly concerning the issues highlighted in the provided text.\n\n"
    f"5. Inquire about production challenges, competitor analysis, and growth-oriented goals, with an emphasis on mitigating risks related to contamination and chemical residues, drawing on the information provided about specific hazards.\n\n"
    f"6. Gather detailed insights into sustainable practices, including resource efficiency, waste management, and the reduction of harmful chemical usage, with reference to the specific chemicals and practices mentioned in the text.\n\n"
    f"7. Capture information regarding the future of company and its commitments to health and sustainability improvements, particularly in relation to minimizing or eliminating the use of harmful substances identified in the provided text.\n\n"
    f"8. Generate Questions in form such that my client has to answer about his product.\n"
    f"9. Don't Use symbols of currency. Rather use Name of Currency in response.\n"
    f"10. Don't use some special symbols(like smart apostrophe) that cannot be encoded using codec.\n"
    f"11. Use good Introduction and conclusion.\n"
    f"12. Question should be of top Quality.\n"
    f"13. Include scientific, geographic, different stages of production process with more emphasis while generating questions.\n"
    f"14. Questions should be detailed and interesting to answer.\n"
    f"15. **Put all Questions of similar topic at once under a sub-heading and do not include question topic separately with each questions.**\n"
    f"16. Generate questions on different processes of making the product.Ex: Different stages of crop production,processing in industries etc.\n"
    f"17. Also add the topic along with the generated questions( examples like: Health Attribute, production type, sustainability, marketing etc.)\n"
    f"18. Questions should be in minimum 40 words and very descriptive and explained.\n\n"
    f"19. Generate questions in such a way that We are assuming that the producers is all clear, but we are going to question them deeply to remove all doubt. Don't start questions with validatory terms.\n"
     
    f"**Follow the questionaire Structure given below:**"
    f"{Text}"
    )           
                st.write(prompt)
                def generate_docx(questions):
                   """
                   Generates a DOCX file with questions and inputs and saves it to disk.
                   """
                   doc = Document()
                   doc.add_heading("Generated Questionnaire", level=1)
    
                   doc.add_heading("Generated Questions:", level=2)
                   for question in questions:
                      doc.add_paragraph(f"- {question}")
    
            
                   doc_path = "questionnaire_report.docx"
                   doc.save(doc_path)
                   return doc_path

                model = genai.GenerativeModel("gemini-2.0-flash-exp")
                response = model.generate_content(prompt)
                questions = response.text.strip().split("\n")

                st.subheader("Generated Questions:")
                for q in questions:
                    st.write(f"- {q}")
                

                pdf_path = generate_pdf(questions)
                docx_path = generate_docx(questions)
              
                with open(pdf_path, "rb") as pdf_file:
                     st.download_button(
                     "Download Full Report as PDF",
                     pdf_file,
                     file_name="questionnaire_report.pdf",
                    )
                with open(docx_path, "rb") as docx_file:
                     st.download_button(
                   "Download Full Report as DOCX",
                   docx_file,
                   file_name="questionnaire_report.docx",
        )
              
            except Exception as e:
                st.error(f"Error generating questions: {e}")
